import docx

def extract_word_content(file_path):
    doc = docx.Document(file_path)
    
    print("=== TEXTO COMPLETO ===")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"[{i+1}] {paragraph.text}")
    
    print("\n=== TABELAS ===")
    for i, table in enumerate(doc.tables):
        print(f"\nTabela {i+1}:")
        for row_idx, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"  Linha {row_idx+1}: {row_text}")
    
    print("\n=== ESTATÍSTICAS ===")
    print(f"Total de parágrafos: {len(doc.paragraphs)}")
    print(f"Total de tabelas: {len(doc.tables)}")
    
    # Verificar se há imagens
    print(f"Imagens encontradas: {len(doc.inline_shapes)}")

if __name__ == "__main__":
    extract_word_content("exemplo_teste/Checklists SGB (1).docx") 